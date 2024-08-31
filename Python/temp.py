from docx import Document

# Create a new Document
doc = Document()

# Add the title
doc.add_heading('Architectural Differences Between RTX and GTX GPUs', 0)

# Add the Keywords section
doc.add_heading('Keywords:', level=1)
keywords = [
    "RTX GPUs",
    "GTX GPUs",
    "Real-Time Ray Tracing",
    "Low-Light Ray Calculation",
    "Gaming Environments",
    "Visual Realism",
]
for keyword in keywords:
    doc.add_paragraph(keyword, style='List Bullet')

# Add the Introduction section
doc.add_heading('Introduction:', level=1)
introduction = (
    "This research paper explores the architectural differences between RTX and GTX GPUs, "
    "with a focus on the impact of real-time ray tracing and low-light ray calculation in gaming environments. "
    "By comparing the two GPU architectures, we aim to understand how these technologies influence visual realism "
    "and performance in gaming. The study delves into the technical aspects of ray tracing, discussing how it is "
    "implemented in RTX GPUs and its absence in GTX GPUs, and how these differences affect the overall gaming experience."
)
doc.add_paragraph(introduction)

# Add the Methodology section
doc.add_heading('Methodology:', level=1)
methodology = (
    "The methodology involves a comparative analysis of the architectures of RTX and GTX GPUs, "
    "focusing on their capabilities in handling ray tracing and low-light scenarios in gaming. "
    "Benchmark tests will be conducted to evaluate the performance differences between these GPUs in various gaming environments. "
    "Data will be collected on frame rates, visual fidelity, and computational overhead to assess the effectiveness of real-time ray tracing. "
    "The results will be analyzed using statistical methods to determine the impact of these architectural differences."
)
doc.add_paragraph(methodology)

# Add the References section
doc.add_heading('References:', level=1)
references = (
    "References will be listed here once the research is completed and sources are cited."
)
doc.add_paragraph(references)

# Save the document
doc.save('Research_Paper_RTX_vs_GTX_GPUs.docx')
